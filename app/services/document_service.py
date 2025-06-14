# app/services/document_service.py
import uuid
import hashlib
import os
import asyncio
from pathlib import Path
from sqlalchemy.orm import Session
from fastapi import UploadFile
from datetime import datetime
import logging

from app.models import database_models as db_m
from app.core.graph import execute_cspe_analysis
from app.config import settings

logger = logging.getLogger(__name__)

# Créer le dossier d'upload s'il n'existe pas
UPLOAD_DIRECTORY = Path(settings.UPLOAD_DIR)
UPLOAD_DIRECTORY.mkdir(exist_ok=True)

class DocumentService:
    """Service pour la gestion des documents et analyses"""
    
    def __init__(self, db: Session):
        self.db = db
    
    async def receive_document(self, file: UploadFile, user_id: str) -> db_m.Document:
        """Reçoit et sauvegarde un document"""
        logger.info(f"📄 Réception du document: {file.filename}")
        
        try:
            # Lire le contenu du fichier
            contents = await file.read()
            content_hash = hashlib.sha256(contents).hexdigest()
            
            # Vérifier si le document existe déjà
            existing_doc = self.db.query(db_m.Document).filter(
                db_m.Document.content_hash == content_hash
            ).first()
            
            if existing_doc:
                logger.info(f"📄 Document existant trouvé: {existing_doc.id}")
                return existing_doc
            
            # Générer un nom de fichier unique
            file_extension = Path(file.filename).suffix if file.filename else '.txt'
            unique_filename = f"{content_hash[:16]}{file_extension}"
            file_path = UPLOAD_DIRECTORY / unique_filename
            
            # Sauvegarder le fichier
            with open(file_path, "wb") as f:
                f.write(contents)
            
            # Créer l'enregistrement en base de données
            new_doc = db_m.Document(
                filename=file.filename or "document_sans_nom.txt",
                content_hash=content_hash,
                file_size=len(contents),
                content_type=file.content_type or "text/plain",
                file_path=str(file_path),
                uploaded_by_id=user_id,
                status=db_m.DocumentStatus.PENDING
            )
            
            self.db.add(new_doc)
            self.db.commit()
            self.db.refresh(new_doc)
            
            logger.info(f"✅ Document sauvegardé: {new_doc.id}")
            return new_doc
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"❌ Erreur lors de la réception du document: {e}")
            raise
    
    def read_document_content(self, document: db_m.Document) -> str:
        """Lit le contenu textuel d'un document"""
        logger.info(f"📖 Lecture du contenu: {document.filename}")
        
        try:
            file_path = Path(document.file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"Fichier non trouvé: {file_path}")
            
            # Lire selon le type de fichier
            if document.content_type == "text/plain" or file_path.suffix == ".txt":
                return self._read_text_file(file_path)
            elif file_path.suffix == ".pdf":
                return self._read_pdf_file(file_path)
            elif file_path.suffix in [".docx", ".doc"]:
                return self._read_word_file(file_path)
            else:
                # Essayer de lire comme texte brut
                return self._read_text_file(file_path)
                
        except Exception as e:
            logger.error(f"❌ Erreur lors de la lecture: {e}")
            raise
    
    def _read_text_file(self, file_path: Path) -> str:
        """Lit un fichier texte"""
        encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.info(f"✅ Fichier lu avec encodage {encoding}")
                return content
            except UnicodeDecodeError:
                continue
        
        # Si aucun encodage ne fonctionne, utiliser errors='ignore'
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        logger.warning("⚠️ Fichier lu avec des caractères ignorés")
        return content
    
    def _read_pdf_file(self, file_path: Path) -> str:
        """Lit un fichier PDF"""
        try:
            import PyMuPDF as fitz
            
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
            
            doc.close()
            logger.info(f"✅ PDF lu: {doc.page_count} pages")
            return text
            
        except ImportError:
            raise ImportError("PyMuPDF requis pour lire les fichiers PDF")
        except Exception as e:
            logger.error(f"❌ Erreur lecture PDF: {e}")
            raise
    
    def _read_word_file(self, file_path: Path) -> str:
        """Lit un fichier Word"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            logger.info("✅ Document Word lu")
            return text
            
        except ImportError:
            raise ImportError("python-docx requis pour lire les fichiers Word")
        except Exception as e:
            logger.error(f"❌ Erreur lecture Word: {e}")
            raise
    
    async def process_document_analysis(self, document_id: str) -> None:
        """Lance l'analyse complète d'un document avec LangGraph"""
        logger.info(f"🤖 Début de l'analyse: {document_id}")
        
        try:
            # Récupérer le document
            document = self.db.query(db_m.Document).filter(
                db_m.Document.id == document_id
            ).first()
            
            if not document:
                raise ValueError(f"Document non trouvé: {document_id}")
            
            # Mettre à jour le statut
            document.status = db_m.DocumentStatus.PROCESSING
            self.db.commit()
            
            # Lire le contenu
            content = self.read_document_content(document)
            
            if not content.strip():
                raise ValueError("Le document est vide ou illisible")
            
            logger.info(f"📄 Contenu lu: {len(content)} caractères")
            
            # Exécuter l'analyse avec LangGraph
            start_time = datetime.utcnow()
            analysis_result = await execute_cspe_analysis(document_id, content)
            end_time = datetime.utcnow()
            
            processing_time_ms = int((end_time - start_time).total_seconds() * 1000)
            
            # Déterminer le résultat de classification
            final_classification = analysis_result.get("final_classification")
            if final_classification == "RECEVABLE":
                classification_result = db_m.ClassificationResult.RECEVABLE
            elif final_classification == "IRRECEVABLE":
                classification_result = db_m.ClassificationResult.IRRECEVABLE
            else:
                classification_result = None  # Nécessite une révision
            
            # Sauvegarder la classification
            classification = db_m.Classification(
                document_id=document_id,
                result=classification_result,
                justification=analysis_result.get("final_justification", ""),
                confidence_score=float(analysis_result.get("final_confidence", 0.0)),
                analysis_steps=analysis_result,
                processing_time_ms=processing_time_ms,
                model_version=settings.LLM_MODEL
            )
            
            self.db.add(classification)
            
            # Mettre à jour le statut du document
            if analysis_result.get("is_review_required", True) or classification_result is None:
                document.status = db_m.DocumentStatus.NEEDS_REVIEW
            else:
                document.status = db_m.DocumentStatus.COMPLETED
            
            self.db.commit()
            
            logger.info(f"✅ Analyse terminée: {final_classification} (confiance: {analysis_result.get('final_confidence', 0):.1%})")
            
        except Exception as e:
            # En cas d'erreur, marquer le document comme erreur
            try:
                document = self.db.query(db_m.Document).filter(
                    db_m.Document.id == document_id
                ).first()
                if document:
                    document.status = db_m.DocumentStatus.ERROR
                    self.db.commit()
            except:
                pass
            
            logger.error(f"❌ Erreur lors de l'analyse: {e}")
            raise
    
    def get_document_with_analysis(self, document_id: str) -> dict:
        """Récupère un document avec son analyse"""
        logger.info(f"📊 Récupération analyse: {document_id}")
        
        try:
            # Récupérer le document
            document = self.db.query(db_m.Document).filter(
                db_m.Document.id == document_id
            ).first()
            
            if not document:
                raise ValueError(f"Document non trouvé: {document_id}")
            
            # Préparer la réponse de base
            result = {
                "id": document.id,
                "filename": document.filename,
                "upload_date": document.upload_date.isoformat(),
                "status": document.status.value,
                "file_size": document.file_size,
                "content_type": document.content_type,
                "classification": None
            }
            
            # Ajouter la classification si elle existe
            if document.classification:
                classification = document.classification
                result["classification"] = {
                    "id": classification.id,
                    "result": classification.result.value if classification.result else None,
                    "justification": classification.justification,
                    "confidence_score": float(classification.confidence_score) if classification.confidence_score else 0.0,
                    "processing_time_ms": classification.processing_time_ms,
                    "model_version": classification.model_version,
                    "created_at": classification.created_at.isoformat(),
                    
                    # Ajouter les détails de l'analyse
                    "final_classification": classification.analysis_steps.get("final_classification") if classification.analysis_steps else None,
                    "final_justification": classification.analysis_steps.get("final_justification") if classification.analysis_steps else None,
                    "final_confidence": classification.analysis_steps.get("final_confidence") if classification.analysis_steps else None,
                    "is_review_required": classification.analysis_steps.get("is_review_required") if classification.analysis_steps else None,
                    "analysis_summary": classification.analysis_steps.get("analysis_summary") if classification.analysis_steps else None
                }
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Erreur lors de la récupération: {e}")
            raise
    
    def delete_document(self, document_id: str) -> bool:
        """Supprime un document et son fichier"""
        logger.info(f"🗑️ Suppression document: {document_id}")
        
        try:
            document = self.db.query(db_m.Document).filter(
                db_m.Document.id == document_id
            ).first()
            
            if not document:
                return False
            
            # Supprimer le fichier physique
            try:
                file_path = Path(document.file_path)
                if file_path.exists():
                    file_path.unlink()
                    logger.info(f"🗑️ Fichier supprimé: {file_path}")
            except Exception as e:
                logger.warning(f"⚠️ Impossible de supprimer le fichier: {e}")
            
            # Supprimer de la base de données (cascade sur classification)
            self.db.delete(document)
            self.db.commit()
            
            logger.info(f"✅ Document supprimé: {document_id}")
            return True
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"❌ Erreur lors de la suppression: {e}")
            return False

# ===== FONCTIONS UTILITAIRES POUR COMPATIBILITÉ =====

async def receive_document(db: Session, file: UploadFile, user_id: str) -> db_m.Document:
    """Fonction de compatibilité avec l'ancien code"""
    service = DocumentService(db)
    return await service.receive_document(file, user_id)

async def process_document_analysis(document_id: str):
    """Fonction de compatibilité avec l'ancien code"""
    from app.database import SessionLocal
    
    db = SessionLocal()
    try:
        service = DocumentService(db)
        await service.process_document_analysis(document_id)
    finally:
        db.close()